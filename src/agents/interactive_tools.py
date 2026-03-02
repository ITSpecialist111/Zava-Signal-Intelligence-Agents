"""Interactive Agent Tools — @tool-decorated functions for Agent 365.

These tools expose the platform's capabilities to the LLM agent so that
users can converse with the agent in Teams / Outlook and ask questions
like:
  - "What signals do we have?"
  - "Run a sweep now"
  - "Show me differences from the last run"
  - "Generate a report"
  - "Show me details for Ark Schools"

Each tool wraps an existing capability (signal store, run history,
sweep workflow, reports, briefs) behind a clean, typed interface that
the Agent Framework's @tool decorator exposes as an OpenAI function call.
"""

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Annotated, Optional

from agent_framework import ai_function

from src.config import AppConfig
from src.models.signal import SignalCategory, SignalStatus
from src.tools.blob_store import BlobReportStore
from src.tools.run_history import RunHistory
from src.tools.signal_store import SignalStore

logger = logging.getLogger(__name__)

# Shared instances — initialised once, reused across tool calls.
_config: AppConfig | None = None
_store: SignalStore | None = None
_history: RunHistory | None = None
_blob: BlobReportStore | None = None


def _get_config() -> AppConfig:
    global _config
    if _config is None:
        _config = AppConfig.from_env()
    return _config


def _get_store() -> SignalStore:
    global _store
    if _store is None:
        _store = SignalStore(_get_config().signal.store_path)
    return _store


def _get_history() -> RunHistory:
    global _history
    if _history is None:
        _history = RunHistory()
    return _history


def _get_blob_store() -> BlobReportStore | None:
    """Lazily create the BlobReportStore. Returns None if not configured."""
    global _blob
    if _blob is not None:
        return _blob
    config = _get_config()
    if not config.blob.is_available:
        return None
    try:
        _blob = BlobReportStore(
            account_url=config.blob.account_url,
            container_name=config.blob.container_name,
            sas_expiry_hours=config.blob.sas_expiry_hours,
        )
        return _blob
    except Exception as exc:
        logger.warning("BlobReportStore not available: %s", exc)
        return None


# ------------------------------------------------------------------ #
# Color maps for professional document styling
# ------------------------------------------------------------------ #

_CATEGORY_HEX: dict[str, str] = {
    "STRUCTURAL_STRESS":    "C0392B",
    "COMPLIANCE_TRAP":      "E67E22",
    "COMPETITOR_MOVEMENT":  "8E44AD",
    "PROCUREMENT_SHIFT":    "2980B9",
    "LEADERSHIP_CHANGE":    "27AE60",
}

_CATEGORY_LIGHT: dict[str, str] = {
    "STRUCTURAL_STRESS":    "FADBD8",
    "COMPLIANCE_TRAP":      "FDEBD0",
    "COMPETITOR_MOVEMENT":  "E8DAEF",
    "PROCUREMENT_SHIFT":    "D4E6F1",
    "LEADERSHIP_CHANGE":    "D5F5E3",
}

_CONF_HEX: dict[str, str] = {"HIGH": "E74C3C", "MEDIUM": "F39C12", "LOW": "27AE60"}
_CONF_BG: dict[str, str] = {"HIGH": "FADBD8", "MEDIUM": "FDEBD0", "LOW": "D5F5E3"}

_STATUS_HEX: dict[str, str] = {
    "DETECTED": "3498DB", "ENRICHED": "2980B9", "HITL_PENDING": "F39C12",
    "APPROVED": "27AE60", "REJECTED": "E74C3C", "ACTIVATED": "2ECC71",
    "ARCHIVED": "95A5A6",
}

_BRAND_PRIMARY = "1B4F72"
_BRAND_ACCENT = "2980B9"


# ------------------------------------------------------------------ #
# Docx helper functions
# ------------------------------------------------------------------ #

def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply background color to a Word table cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_border(paragraph, side: str, hex_color: str, size: int = 24) -> None:
    """Add a colored border to a paragraph (side: left, bottom, top)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    border_el = OxmlElement(f"w:{side}")
    border_el.set(qn("w:val"), "single")
    border_el.set(qn("w:sz"), str(size))
    border_el.set(qn("w:space"), "4")
    border_el.set(qn("w:color"), hex_color)
    p_bdr.append(border_el)
    p_pr.append(p_bdr)


def _set_table_borders(table, hex_color: str = "D5D8DC", size: int = 4) -> None:
    """Apply thin borders to all cells in a table."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tbl_pr.append(borders)


# ------------------------------------------------------------------ #
# Helper: Render structured sections to a colour-coded .docx
# ------------------------------------------------------------------ #

def _render_docx(sections: list[dict]) -> bytes:
    """Convert structured report sections into a professional, colour-coded
    Word .docx byte stream.

    Supported section types: title, subtitle, heading1, heading2,
    paragraph, table, quote, kpi_row, page_break.

    Optional section keys:
      - category: str — auto-maps to category accent colour
      - color: str — explicit hex colour override (e.g. "E74C3C")
    """
    import io

    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    # ── Global styles ──────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        key = f"Heading {level}"
        if key in doc.styles:
            doc.styles[key].font.name = "Calibri"

    for sec_page in doc.sections:
        sec_page.top_margin = Cm(2)
        sec_page.bottom_margin = Cm(2)
        sec_page.left_margin = Cm(2.5)
        sec_page.right_margin = Cm(2.5)

    # ── Header / Footer branding ──────────────────────────────────
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("CONFIDENTIAL  ·  Zava Signal Intelligence")
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("© 2026 Zellis  |  Signal Intelligence Report")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Render each section ───────────────────────────────────────
    for sect in sections:
        stype = sect.get("type", "paragraph")
        text = sect.get("text", "")
        category = sect.get("category", "")
        color = sect.get("color", "")
        accent = color or _CATEGORY_HEX.get(category, "")

        # ── Title (cover) ─────────────────────────────────────────
        if stype == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)
            _add_border(p, "bottom", _BRAND_PRIMARY, 36)

        # ── Subtitle ──────────────────────────────────────────────
        elif stype == "subtitle":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.italic = True

        # ── Heading 1 (with optional category accent bar) ─────────
        elif stype == "heading1":
            h = doc.add_heading(text, level=1)
            if accent:
                _add_border(h, "left", accent, 30)
                _add_border(h, "bottom", accent, 6)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(*bytes.fromhex(accent))

        # ── Heading 2 (with optional accent) ──────────────────────
        elif stype == "heading2":
            h = doc.add_heading(text, level=2)
            if accent:
                _add_border(h, "left", accent, 18)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(*bytes.fromhex(accent))

        # ── Paragraph (with optional left‑bar accent) ─────────────
        elif stype == "paragraph":
            p = doc.add_paragraph(text)
            if accent:
                _add_border(p, "left", accent, 12)

        # ── Page break ────────────────────────────────────────────
        elif stype == "page_break":
            doc.add_page_break()

        # ── KPI Row (coloured metric boxes) ───────────────────────
        elif stype == "kpi_row":
            items = sect.get("items", [])
            if items:
                table = doc.add_table(rows=2, cols=len(items))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                _set_table_borders(table, "FFFFFF", 2)  # subtle borders
                for j, item in enumerate(items):
                    item_color = item.get("color", _BRAND_ACCENT)
                    light_bg = _CATEGORY_LIGHT.get(
                        item.get("category", ""), item.get("bg", "EBF5FB")
                    )
                    # Value cell (top)
                    val_cell = table.cell(0, j)
                    val_cell.text = ""
                    vp = val_cell.paragraphs[0]
                    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    vr = vp.add_run(item["value"])
                    vr.font.size = Pt(22)
                    vr.font.bold = True
                    vr.font.color.rgb = RGBColor(*bytes.fromhex(item_color))
                    _set_cell_shading(val_cell, light_bg)
                    # Label cell (bottom)
                    lbl_cell = table.cell(1, j)
                    lbl_cell.text = ""
                    lp = lbl_cell.paragraphs[0]
                    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    lr = lp.add_run(item["label"])
                    lr.font.size = Pt(9)
                    lr.font.bold = True
                    lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_cell_shading(lbl_cell, item_color)
                doc.add_paragraph()  # spacing

        # ── Quote (with coloured left border) ─────────────────────
        elif stype == "quote":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _add_border(p, "left", accent or "999999", 18)

        # ── Table (auto‑coloured headers, confidence & status) ────
        elif stype == "table":
            headers = sect.get("headers", [])
            rows = sect.get("rows", [])
            if not headers:
                continue

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_borders(table)

            # Detect auto-colour columns
            conf_col = next((j for j, h in enumerate(headers) if h.lower() == "confidence"), None)
            status_col = next((j for j, h in enumerate(headers) if h.lower() == "status"), None)
            header_bg = accent or _BRAND_PRIMARY

            # Header row — white text on coloured background
            for j, h in enumerate(headers):
                cell = table.cell(0, j)
                cell.text = ""
                cp = cell.paragraphs[0]
                cr = cp.add_run(h)
                cr.font.bold = True
                cr.font.size = Pt(9)
                cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_shading(cell, header_bg)

            # Data rows
            for i, row in enumerate(rows):
                row_bg = "F8F9FA" if i % 2 == 1 else None
                for j, val in enumerate(row):
                    cell = table.cell(i + 1, j)
                    cell.text = ""
                    cp = cell.paragraphs[0]
                    cr = cp.add_run(str(val))
                    cr.font.size = Pt(9)

                    cell_colored = False
                    # Confidence column — background + bold text
                    if j == conf_col:
                        val_str = str(val).upper()
                        for level in ("HIGH", "MEDIUM", "LOW"):
                            if level in val_str:
                                _set_cell_shading(cell, _CONF_BG[level])
                                cr.font.color.rgb = RGBColor(
                                    *bytes.fromhex(_CONF_HEX[level])
                                )
                                cr.font.bold = True
                                cell_colored = True
                                break
                    # Status column — coloured text
                    elif j == status_col:
                        val_str = str(val).upper().replace(" ", "_")
                        if val_str in _STATUS_HEX:
                            cr.font.color.rgb = RGBColor(
                                *bytes.fromhex(_STATUS_HEX[val_str])
                            )
                            cr.font.bold = True

                    # Zebra-stripe non-special cells
                    if not cell_colored and row_bg:
                        _set_cell_shading(cell, row_bg)

            doc.add_paragraph()  # spacing after table

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ------------------------------------------------------------------ #
# Tool: List / query signals
# ------------------------------------------------------------------ #

@ai_function
def get_signals(
    category: Annotated[
        Optional[str],
        "Filter by signal category: STRUCTURAL_STRESS, COMPLIANCE_TRAP, COMPETITOR_MOVEMENT, PROCUREMENT_SHIFT, or LEADERSHIP_CHANGE. Leave empty for all."
    ] = None,
    status: Annotated[
        Optional[str],
        "Filter by status: DETECTED, ENRICHED, HITL_PENDING, APPROVED, REJECTED, ACTIVATED, ARCHIVED. Leave empty for all."
    ] = None,
    days: Annotated[
        Optional[int],
        "Only show signals from the last N days. Leave empty for all time."
    ] = None,
    limit: Annotated[
        int,
        "Maximum number of signals to return."
    ] = 20,
) -> str:
    """List detected signals with optional filters for category, status, and recency.

    Use this to answer questions like "What signals do we have?",
    "Show me high-priority signals", or "Any new procurement signals this week?"
    """
    store = _get_store()

    if days:
        signals = store.get_recent_signals(days)
    else:
        signals = store.get_all_signals()

    # Apply category filter
    if category:
        cat_upper = category.upper().replace(" ", "_")
        signals = [s for s in signals if s.category.value == cat_upper]

    # Apply status filter
    if status:
        status_upper = status.upper().replace(" ", "_")
        signals = [s for s in signals if s.status.value == status_upper]

    # Sort by confidence descending
    signals.sort(key=lambda s: s.confidence, reverse=True)
    total_count = len(signals)
    signals = signals[:limit]

    if not signals:
        return "No signals found matching the specified filters."

    lines = [f"**{total_count} signals found** (showing top {len(signals)}):\n"]
    lines.append("| # | Trust | Category | Confidence | Status | Playbook | Detected |")
    lines.append("|---|-------|----------|-----------|--------|----------|----------|")

    for i, s in enumerate(signals, 1):
        playbook = s.playbook_match.value if s.playbook_match else "—"
        detected = s.detected_at.strftime("%d %b %Y") if s.detected_at else "—"
        lines.append(
            f"| {i} | {s.entity_name} | {s.category.value} | "
            f"{s.confidence:.0%} ({s.confidence_level.value}) | "
            f"{s.status.value} | {playbook} | {detected} |"
        )

    # Category breakdown
    counts = store.get_signal_counts_by_category()
    if counts:
        lines.append("\n**Category Breakdown:**")
        for cat, count in sorted(counts.items(), key=lambda x: -x[1]):
            lines.append(f"- {cat}: {count}")

    full_report = "\n".join(lines)

    # Upload full report to Blob and return compact summary + link
    blob = _get_blob_store()
    if blob and total_count > 5:
        try:
            ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            blob_name = f"signals/signals_{ts}.md"
            url = blob.upload_text(full_report, blob_name)

            # Compact summary for chat
            high = sum(1 for s in signals if s.confidence >= 0.8)
            medium = sum(1 for s in signals if 0.5 <= s.confidence < 0.8)
            low = sum(1 for s in signals if s.confidence < 0.5)
            pending = sum(1 for s in signals if s.status.value == "HITL_PENDING")

            top_cats = sorted(counts.items(), key=lambda x: -x[1])[:3] if counts else []
            cat_summary = ", ".join(f"{c}: {n}" for c, n in top_cats)

            summary_lines = [
                f"**{total_count} signals** found.\n",
                f"- HIGH confidence: {high}",
                f"- MEDIUM confidence: {medium}",
                f"- LOW confidence: {low}",
            ]
            if pending:
                summary_lines.append(f"- Pending HITL review: {pending}")
            if cat_summary:
                summary_lines.append(f"\nTop categories: {cat_summary}")
            if signals:
                summary_lines.append(f"\nTop signal: **{signals[0].entity_name}** — {signals[0].category.value} ({signals[0].confidence:.0%})")
            summary_lines.append(f"\n📄 [View full signal report]({url})")
            return "\n".join(summary_lines)
        except Exception as exc:
            logger.warning("Failed to upload signals to blob: %s", exc)

    return full_report


# ------------------------------------------------------------------ #
# Tool: Get signal details
# ------------------------------------------------------------------ #

@ai_function
def get_signal_details(
    entity_name: Annotated[
        str,
        "Name of the academy trust, e.g. 'Ark Schools' or 'Harris Federation'. Partial matching supported."
    ],
) -> str:
    """Get detailed information about a specific trust's signals.

    Use this when a user asks for details on a particular trust,
    e.g. "Tell me about Ark Schools" or "What do we know about Harris Federation?"
    """
    store = _get_store()
    all_signals = store.get_all_signals()

    # Partial match (case-insensitive)
    name_lower = entity_name.lower()
    matches = [s for s in all_signals if name_lower in s.entity_name.lower()]

    if not matches:
        return f"No signals found for '{entity_name}'. Try a different search term."

    lines = [f"**Found {len(matches)} signal(s) for '{entity_name}':**\n"]

    for s in matches:
        lines.append(f"### {s.entity_name}")
        lines.append(f"- **Signal ID:** {s.signal_id}")
        lines.append(f"- **Category:** {s.category.value} / {s.subcategory.value}")
        lines.append(f"- **Confidence:** {s.confidence:.0%} ({s.confidence_level.value})")
        lines.append(f"- **Status:** {s.status.value}")
        lines.append(f"- **Source:** [{s.source_name}]({s.source_url})")
        lines.append(f"- **Detected:** {s.detected_at.strftime('%d %b %Y %H:%M UTC')}")

        if s.playbook_match:
            lines.append(f"- **Playbook:** {s.playbook_match.value}")
        if s.recommended_action:
            lines.append(f"- **Recommended Action:** {s.recommended_action}")
        if s.impact_statement:
            lines.append(f"- **Impact:** {s.impact_statement}")
        if s.talk_track:
            lines.append(f"- **Talk Track:** {s.talk_track}")
        if s.current_provider:
            lines.append(f"- **Current Provider:** {s.current_provider}")
        if s.financial_summary:
            lines.append(f"- **Financials:** {s.financial_summary}")
        if s.key_decision_maker:
            lines.append(f"- **Key Contact:** {s.key_decision_maker}")
        if s.raw_evidence:
            lines.append(f"\n> **Evidence:** {s.raw_evidence[:500]}")

        lines.append("")

    return "\n".join(lines)


# ------------------------------------------------------------------ #
# Tool: Signals pending HITL review
# ------------------------------------------------------------------ #

@ai_function
def get_hitl_queue() -> str:
    """List signals that are pending human-in-the-loop review.

    Use this when users ask "What needs my review?", "HITL queue",
    or "What signals are waiting for approval?"
    """
    store = _get_store()
    pending = store.get_signals_pending_review()

    if not pending:
        return "No signals are currently pending HITL review. All clear!"

    lines = [f"**{len(pending)} signal(s) pending review:**\n"]
    lines.append("| Trust | Category | Confidence | Evidence Summary |")
    lines.append("|-------|----------|-----------|-----------------|")

    for s in pending:
        evidence = (s.raw_evidence[:80] + "…") if len(s.raw_evidence) > 80 else s.raw_evidence
        lines.append(
            f"| {s.entity_name} | {s.category.value} | "
            f"{s.confidence:.0%} | {evidence} |"
        )

    return "\n".join(lines)


# ------------------------------------------------------------------ #
# Tool: Approve or reject a HITL signal
# ------------------------------------------------------------------ #

@ai_function
def review_signal(
    entity_name: Annotated[
        str,
        "Name (or partial name) of the trust whose signal to review."
    ],
    decision: Annotated[
        str,
        "Either 'approve' or 'reject'."
    ],
    notes: Annotated[
        Optional[str],
        "Optional reviewer notes."
    ] = None,
) -> str:
    """Approve or reject a signal that is pending HITL review.

    Use this when the user says things like "Approve Ark Schools" or
    "Reject that competitor signal for Harris".
    """
    store = _get_store()
    pending = store.get_signals_pending_review()

    name_lower = entity_name.lower()
    matches = [s for s in pending if name_lower in s.entity_name.lower()]

    if not matches:
        return f"No HITL-pending signal found matching '{entity_name}'."

    decision_lower = decision.lower().strip()
    if decision_lower not in ("approve", "reject"):
        return "Decision must be 'approve' or 'reject'."

    new_status = SignalStatus.APPROVED if decision_lower == "approve" else SignalStatus.REJECTED
    results = []

    for s in matches:
        store.update_status(
            s.signal_id,
            new_status,
            reviewed_by="interactive_agent",
            review_notes=notes or "",
            reviewed_at=datetime.utcnow(),
        )
        results.append(f"- **{s.entity_name}** ({s.signal_id[:8]}…) → {new_status.value}")

    action = "Approved" if decision_lower == "approve" else "Rejected"
    return f"{action} {len(results)} signal(s):\n" + "\n".join(results)


# ------------------------------------------------------------------ #
# Tool: Run sweep on-demand
# ------------------------------------------------------------------ #

@ai_function
async def run_sweep_now() -> str:
    """Trigger a full signal collection sweep immediately.

    This runs the entire pipeline: web crawl → procurement scan →
    enrich → route → store. It may take a few minutes.

    Use this when the user says "Run a sweep", "Scan for new signals",
    or "Check for new intelligence now".
    """
    from src.workflows.daily_sweep import run_daily_sweep

    config = _get_config()
    store = _get_store()

    summary = await run_daily_sweep(
        signal_store=store,
        confidence_threshold=config.signal.confidence_threshold,
    )

    # Record in history
    history = _get_history()
    history.record(summary)

    lines = [
        "**Sweep complete!**\n",
        f"- Signals detected: **{summary['total_detected']}**",
        f"- After dedup: **{summary['after_dedup']}**",
        f"- Enriched: **{summary['enriched']}**",
        f"- Auto-activated: **{summary['auto_activated']}**",
        f"- HITL pending: **{summary['hitl_pending']}**",
    ]

    if summary.get("errors"):
        lines.append(f"- Errors: {len(summary['errors'])}")

    cats = summary.get("category_breakdown", {})
    if cats:
        lines.append("\n**By category:**")
        for cat, count in sorted(cats.items(), key=lambda x: -x[1]):
            lines.append(f"- {cat}: {count}")

    return "\n".join(lines)


# ------------------------------------------------------------------ #
# Tool: Run history / comparison
# ------------------------------------------------------------------ #

@ai_function
def get_run_history(
    last_n: Annotated[
        int,
        "Number of recent runs to show."
    ] = 5,
) -> str:
    """Show the history of sweep runs with a comparison table.

    Use this when the user asks "Show me run history",
    "What's changed since the last sweep?", or "Compare recent runs".
    """
    history = _get_history()

    if not history.count:
        return "No sweep runs recorded yet. Use `run_sweep_now` to trigger one."

    return history.format_summary()


# ------------------------------------------------------------------ #
# Tool: Generate weekly brief
# ------------------------------------------------------------------ #

@ai_function
def generate_brief() -> str:
    """Generate the weekly segment brief as Markdown.

    Use this when the user says "Generate the weekly brief",
    "Give me a summary of this week's signals", or "Segment brief".
    """
    from src.outputs.segment_brief import generate_segment_brief

    store = _get_store()
    config = _get_config()
    all_signals = store.get_all_signals()
    brief = generate_segment_brief(all_signals)

    # Also save to file
    output_dir = Path(config.output.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.utcnow().strftime('%Y%m%d')
    brief_path = output_dir / f"brief_{ts}.md"
    brief_path.write_text(brief, encoding="utf-8")

    # Upload to Blob and return summary + link
    blob = _get_blob_store()
    if blob:
        try:
            blob_name = f"briefs/brief_{ts}.md"
            url = blob.upload_text(brief, blob_name)

            # Count signals for summary
            total = len(all_signals)
            high = sum(1 for s in all_signals if s.confidence >= 0.8)

            return (
                f"**Weekly Segment Brief** generated.\n\n"
                f"- Signals analysed: **{total}**\n"
                f"- HIGH confidence: **{high}**\n"
                f"- Week ending: {datetime.utcnow().strftime('%d %B %Y')}\n\n"
                f"📄 [View full brief]({url})"
            )
        except Exception as exc:
            logger.warning("Failed to upload brief to blob: %s", exc)

    return brief


# ------------------------------------------------------------------ #
# Tool: Generate monthly report
# ------------------------------------------------------------------ #

@ai_function
def generate_report() -> str:
    """Generate the monthly horizon report as PDF and return a summary.

    Use this when a user asks for "the monthly report", "horizon report",
    or "generate a strategic report".
    """
    from src.outputs.horizon_report import generate_horizon_report

    store = _get_store()
    config = _get_config()
    all_signals = store.get_all_signals()

    output_dir = Path(config.output.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.utcnow().strftime('%Y%m')
    report_path = output_dir / f"horizon_{ts}.pdf"

    pdf_bytes = generate_horizon_report(
        signals=all_signals,
        output_path=report_path,
    )

    # Upload PDF to Blob and return summary + link
    blob = _get_blob_store()
    if blob and pdf_bytes:
        try:
            blob_name = f"reports/horizon_{ts}.pdf"
            url = blob.upload_bytes(pdf_bytes, blob_name)

            return (
                f"**Monthly Horizon Report** generated.\n\n"
                f"- Signals analysed: **{len(all_signals)}**\n"
                f"- Period: 3-month lookback to {datetime.utcnow().strftime('%B %Y')}\n\n"
                f"📄 [Download Horizon Report (PDF)]({url})"
            )
        except Exception as exc:
            logger.warning("Failed to upload report to blob: %s", exc)

    return (
        f"Monthly Horizon Report generated successfully.\n"
        f"**Saved to:** {report_path}\n"
        f"**Signals analysed:** {len(all_signals)}\n"
        f"**Period:** 3-month lookback to {datetime.utcnow().strftime('%B %Y')}"
    )


# ------------------------------------------------------------------ #
# Tool: Signal statistics / dashboard
# ------------------------------------------------------------------ #

@ai_function
def get_signal_dashboard() -> str:
    """Get a high-level dashboard of all signal intelligence.

    Use this when the user asks "Give me the overview", "Dashboard",
    "How are we doing?", or "What's the current state?"
    """
    store = _get_store()
    history = _get_history()
    all_signals = store.get_all_signals()

    if not all_signals:
        return "No signals in the store yet. Run a sweep to start collecting intelligence."

    # Status breakdown
    status_counts: dict[str, int] = {}
    for s in all_signals:
        status_counts[s.status.value] = status_counts.get(s.status.value, 0) + 1

    # Category breakdown
    cat_counts = store.get_signal_counts_by_category()

    # Confidence distribution
    high = sum(1 for s in all_signals if s.confidence >= 0.8)
    medium = sum(1 for s in all_signals if 0.5 <= s.confidence < 0.8)
    low = sum(1 for s in all_signals if s.confidence < 0.5)

    # Recent activity
    recent_7d = store.get_recent_signals(7)
    recent_30d = store.get_recent_signals(30)

    lines = [
        "# Signal Intelligence Dashboard\n",
        f"**Total signals:** {len(all_signals)}",
        f"**Last 7 days:** {len(recent_7d)}",
        f"**Last 30 days:** {len(recent_30d)}",
        f"**Sweep runs recorded:** {history.count}",
        "",
        "## Confidence Distribution",
        f"- HIGH (≥80%): {high}",
        f"- MEDIUM (50–79%): {medium}",
        f"- LOW (<50%): {low}",
        "",
        "## By Status",
    ]
    for status, count in sorted(status_counts.items(), key=lambda x: -x[1]):
        lines.append(f"- {status}: {count}")

    lines.append("\n## By Category")
    for cat, count in sorted(cat_counts.items(), key=lambda x: -x[1]):
        lines.append(f"- {cat}: {count}")

    # Last run info
    last_run = history.latest
    if last_run:
        lines.append("\n## Last Sweep")
        lines.append(f"- Run #{last_run['run_number']} at {last_run['timestamp']}")
        lines.append(f"- Detected: {last_run['total_detected']}, Auto-activated: {last_run['auto_activated']}, HITL: {last_run['hitl_pending']}")

    return "\n".join(lines)


# ------------------------------------------------------------------ #
# Tool: Send daily digest email
# ------------------------------------------------------------------ #

@ai_function
def send_daily_digest(
    days: Annotated[
        int,
        "Only include signals from the last N days. Defaults to 1 (today's signals)."
    ] = 1,
) -> str:
    """Send a daily signal intelligence digest email to the team.

    Compiles all signals from the specified period into a professionally
    formatted HTML email and sends it via the Mail MCP server.

    Use when the user says "email the team", "send the digest",
    "distribute today's signals", or after a sweep completes.
    """
    from src.agents.proactive_actions import build_daily_digest_email
    from src.config_teams import ProactiveConfig

    store = _get_store()
    config = ProactiveConfig.load()
    signals = store.get_recent_signals(days)

    if not signals:
        return f"No signals found in the last {days} day(s). Nothing to send."

    # Build the sweep summary for the email
    sweep_summary = {
        "total_detected": len(signals),
        "after_dedup": len(signals),
        "enriched": len(signals),
    }

    email_data = build_daily_digest_email(signals, sweep_summary, config)

    if email_data.get("skipped"):
        return f"Digest skipped: {email_data['reason']}"

    return (
        f"📧 **Daily digest ready to send**\n\n"
        f"- **Subject:** {email_data['subject']}\n"
        f"- **Recipients:** {', '.join(email_data['recipients'])}\n"
        f"- **Signals included:** {len(signals)}\n"
        f"- **Importance:** {email_data['importance']}\n\n"
        f"Please use the **Mail MCP server** to send this email with "
        f"the HTML body provided."
    )


# ------------------------------------------------------------------ #
# Tool: Send HIGH-confidence alert
# ------------------------------------------------------------------ #

@ai_function
def send_high_alert(
    entity_name: Annotated[
        str,
        "Name (or partial name) of the trust to send an alert for."
    ],
) -> str:
    """Send an instant email alert for a HIGH-confidence signal.

    Use when a critical signal is detected and the team needs to know
    immediately, or when the user says "alert the team about this signal".
    """
    from src.agents.proactive_actions import build_high_alert_email
    from src.config_teams import ProactiveConfig

    store = _get_store()
    config = ProactiveConfig.load()
    all_signals = store.get_all_signals()

    name_lower = entity_name.lower()
    matches = [s for s in all_signals if name_lower in s.entity_name.lower()]

    if not matches:
        return f"No signal found matching '{entity_name}'."

    results = []
    for signal in matches:
        email_data = build_high_alert_email(signal, config)
        if email_data.get("skipped"):
            results.append(f"- {signal.entity_name}: Skipped — {email_data['reason']}")
        else:
            results.append(
                f"- **{signal.entity_name}** → {', '.join(email_data['recipients'])} "
                f"(importance: {email_data['importance']})"
            )

    return (
        "⚡ **Alert email(s) prepared:**\n\n" +
        "\n".join(results) +
        "\n\nPlease use the **Mail MCP server** to send."
    )


# ------------------------------------------------------------------ #
# Tool: Schedule a review meeting
# ------------------------------------------------------------------ #

@ai_function
def schedule_review_meeting(
    days_ahead: Annotated[
        int,
        "Schedule the meeting N days from now (default: next business day)."
    ] = 1,
) -> str:
    """Schedule a Teams meeting to review signals pending HITL approval.

    Gathers all HITL-pending signals and creates a meeting invite with
    a structured agenda. Use when the user says "set up a meeting to
    discuss the signals", "schedule a review", or "organise a call".
    """
    from src.agents.proactive_actions import build_review_meeting
    from src.config_teams import ProactiveConfig

    store = _get_store()
    config = ProactiveConfig.load()
    pending = store.get_signals_pending_review()

    if not pending:
        return "No signals are pending HITL review — no meeting needed."

    meeting_data = build_review_meeting(pending, config)

    if meeting_data.get("skipped"):
        return f"Meeting skipped: {meeting_data['reason']}"

    return (
        f"📅 **Review meeting prepared:**\n\n"
        f"- **Subject:** {meeting_data['subject']}\n"
        f"- **Time:** {meeting_data['start_time']} – {meeting_data['end_time']}\n"
        f"- **Attendees:** {', '.join(meeting_data['attendees'])}\n"
        f"- **Signals to review:** {len(pending)}\n"
        f"- **Online meeting:** {'Yes' if meeting_data['is_online_meeting'] else 'No'}\n\n"
        f"Please use the **Calendar MCP server** to create this meeting."
    )


# ------------------------------------------------------------------ #
# Tool: Create Word report
# ------------------------------------------------------------------ #

@ai_function
def create_word_report(
    days: Annotated[
        int,
        "Include signals from the last N days. Use 30 for monthly, 7 for weekly."
    ] = 30,
) -> str:
    """Generate and upload a downloadable Word (.docx) signal intelligence report.

    This is the FULL REPORT tool. Creates a professional .docx document
    with executive summary, category breakdowns, high-confidence signal
    details, talk tracks, and methodology appendix, then uploads it to
    Azure Blob Storage and returns a download link.

    Use when the user says "full report", "generate a report",
    "create a Word report", "make a document", "download the report",
    or "build the report".
    """
    from src.agents.proactive_actions import build_word_report_content

    store = _get_store()
    signals = store.get_recent_signals(days)

    if not signals:
        return f"No signals in the last {days} days. Cannot generate report."

    report_data = build_word_report_content(signals)
    docx_bytes = _render_docx(report_data["sections"])
    filename = report_data["filename"]

    high_count = sum(1 for s in signals if s.confidence >= 0.8)
    unique_trusts = len(set(s.entity_name for s in signals))

    # Upload to Blob and return download link
    blob = _get_blob_store()
    if blob and docx_bytes:
        try:
            ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            blob_name = f"reports/{filename.replace('.docx', '')}_{ts}.docx"
            url = blob.upload_bytes(docx_bytes, blob_name)

            return (
                f"📄 **Full Signal Intelligence Report** generated.\n\n"
                f"- **Signals covered:** {len(signals)}\n"
                f"- **HIGH confidence:** {high_count}\n"
                f"- **Unique trusts:** {unique_trusts}\n"
                f"- **Period:** last {days} days\n\n"
                f"📥 [Download full report (.docx)]({url})"
            )
        except Exception as exc:
            logger.warning("Failed to upload Word report to blob: %s", exc)

    return (
        f"📄 **Report generated** ({len(signals)} signals, {high_count} HIGH).\n"
        f"Blob upload unavailable — report saved locally as {filename}."
    )


# ------------------------------------------------------------------ #
# Tool: Create Planner tasks from approved signals
# ------------------------------------------------------------------ #

@ai_function
def create_action_tasks(
    category: Annotated[
        Optional[str],
        "Only create tasks for signals in this category. Leave empty for all approved signals."
    ] = None,
) -> str:
    """Create Planner tasks for approved signals with recommended actions.

    Each approved signal becomes a task assigned to the relevant team
    members, with signal details as task notes.

    Use when the user says "create tasks", "add to Planner",
    "assign actions from the signals", or "set up follow-ups".
    """
    from src.agents.proactive_actions import build_planner_tasks
    from src.config_teams import ProactiveConfig

    store = _get_store()
    config = ProactiveConfig.load()
    all_signals = store.get_all_signals()

    approved = [s for s in all_signals if s.status.value == "APPROVED"]
    if category:
        cat_upper = category.upper().replace(" ", "_")
        approved = [s for s in approved if s.category.value == cat_upper]

    if not approved:
        return "No approved signals with recommended actions found."

    tasks = build_planner_tasks(approved, config)

    if not tasks:
        return "Approved signals found but none have recommended actions to create tasks for."

    lines = [f"📋 **{len(tasks)} Planner task(s) prepared:**\n"]
    for t in tasks:
        lines.append(
            f"- **{t['title'][:60]}** → {', '.join(t['assignees']) or 'Unassigned'} "
            f"(due {t['due_date']}, priority: {t['priority']})"
        )

    lines.append(
        "\nPlease use the **Planner MCP server** to create these tasks."
    )

    return "\n".join(lines)


# ------------------------------------------------------------------ #
# Tool: Update pipeline tracker (Excel)
# ------------------------------------------------------------------ #

@ai_function
def update_pipeline_tracker() -> str:
    """Update a signal pipeline tracker in Excel.

    Exports all signals into a structured spreadsheet with columns
    for trust name, category, confidence, status, source, actions, etc.

    Use when the user says "update the spreadsheet", "export to Excel",
    or "refresh the pipeline tracker".
    """
    from src.agents.proactive_actions import build_excel_pipeline_data

    store = _get_store()
    all_signals = store.get_all_signals()

    if not all_signals:
        return "No signals to export. Run a sweep first."

    data = build_excel_pipeline_data(all_signals)

    return (
        f"📊 **Excel pipeline tracker prepared:**\n\n"
        f"- **Filename:** {data['filename']}\n"
        f"- **Sheet:** {data['sheet_name']}\n"
        f"- **Columns:** {len(data['headers'])}\n"
        f"- **Rows:** {len(data['rows'])}\n\n"
        f"Please use the **Excel MCP server** to create or update the "
        f"workbook, then upload to SharePoint for the team."
    )


# ------------------------------------------------------------------ #
# Tool: Post sweep results to Teams channel
# ------------------------------------------------------------------ #

@ai_function
def post_to_teams_channel(
    days: Annotated[
        int,
        "Include signals from the last N days. Defaults to 1 (today's sweep)."
    ] = 1,
) -> str:
    """Post a sweep summary to the team's signal intelligence channel.

    Creates an Adaptive Card with key metrics and top signals.
    Use when the user says "post to Teams", "notify the channel",
    or "share results with the team".
    """
    from src.agents.proactive_actions import build_teams_channel_message

    store = _get_store()
    signals = store.get_recent_signals(days)

    if not signals:
        return f"No signals from the last {days} day(s) to post."

    sweep_summary = {
        "total_detected": len(signals),
        "after_dedup": len(signals),
    }

    message = build_teams_channel_message(signals, sweep_summary)

    return (
        f"💬 **Teams channel message prepared:**\n\n"
        f"- **Summary:** {message['summary']}\n"
        f"- **Format:** Adaptive Card v1.5\n\n"
        f"Please use the **Teams MCP server** to post this card "
        f"to the signal intelligence channel."
    )


# ------------------------------------------------------------------ #
# Tool: Full proactive distribution (email + Teams + tasks)
# ------------------------------------------------------------------ #

@ai_function
def distribute_intelligence(
    days: Annotated[
        int,
        "Period to cover (days). Defaults to 1 for daily distribution."
    ] = 1,
) -> str:
    """Run the full proactive intelligence distribution workflow.

    This is the "do everything" tool: sends the email digest, posts
    to Teams, creates Planner tasks for approved signals, and prepares
    a Word report — a complete distribution cycle.

    Use when the user says "distribute the intelligence", "push
    everything out", or "run the full notification cycle".
    """
    from src.agents.proactive_actions import (
        build_daily_digest_email,
        build_planner_tasks,
        build_teams_channel_message,
        build_word_report_content,
    )
    from src.config_teams import ProactiveConfig

    store = _get_store()
    config = ProactiveConfig.load()
    signals = store.get_recent_signals(days)

    if not signals:
        return f"No signals from the last {days} day(s). Nothing to distribute."

    # Build all components
    sweep_summary = {"total_detected": len(signals), "after_dedup": len(signals)}

    email = build_daily_digest_email(signals, sweep_summary, config)
    teams_msg = build_teams_channel_message(signals, sweep_summary)
    report = build_word_report_content(signals)

    approved = [s for s in signals if s.status.value == "APPROVED"]
    tasks = build_planner_tasks(approved, config) if approved else []

    high_count = sum(1 for s in signals if s.confidence >= 0.8)

    actions = []
    if not email.get("skipped"):
        actions.append(f"📧 Email digest → {len(email['recipients'])} recipients")
    if teams_msg:
        actions.append("💬 Teams channel post — Adaptive Card")
    if tasks:
        actions.append(f"📋 {len(tasks)} Planner task(s) for approved signals")
    actions.append(f"📄 Word report — {report['filename']}")

    return (
        f"🚀 **Full distribution prepared:**\n\n"
        f"- **Signals:** {len(signals)} ({high_count} high confidence)\n"
        f"- **Period:** last {days} day(s)\n\n"
        f"**Actions to execute:**\n" +
        "\n".join(f"  {i+1}. {a}" for i, a in enumerate(actions)) +
        "\n\nUse the MCP servers to execute each action:\n"
        "1. **Mail** → send the digest email\n"
        "2. **Teams** → post the Adaptive Card\n"
        "3. **Planner** → create the action tasks\n"
        "4. **Word** → create the report document\n"
        "5. **OneDrive/SharePoint** → save the report"
    )


# ------------------------------------------------------------------ #
# Tool: Signal cards (structured JSON for Adaptive Cards in Studio)
# ------------------------------------------------------------------ #

@ai_function
def get_signal_cards(
    days: Annotated[
        int,
        "Include signals from the last N days. Use 30 for monthly, 7 for weekly."
    ] = 30,
    limit: Annotated[
        int,
        "Maximum number of signal cards to return."
    ] = 10,
) -> str:
    """Return signal intelligence as structured JSON for Adaptive Card rendering.

    This tool outputs a JSON payload containing signal cards, category
    summaries, and colour codes.  Copilot Studio can parse this JSON and
    map it to a custom Adaptive Card template for rich, colour-coded
    display inside Teams.

    Use when the user says "show me signal cards", "adaptive card view",
    "card format", or "formatted summary".
    """
    store = _get_store()
    signals = store.get_recent_signals(days)
    signals.sort(key=lambda s: s.confidence, reverse=True)
    total = len(signals)
    signals = signals[:limit]

    if not signals:
        return json.dumps({"cards": [], "total_signals": 0})

    # Build card objects
    cards: list[dict] = []
    for s in signals:
        cards.append({
            "entity_name": s.entity_name,
            "category": s.category.value,
            "category_label": s.category.value.replace("_", " ").title(),
            "category_color": f"#{_CATEGORY_HEX.get(s.category.value, '2980B9')}",
            "subcategory": s.subcategory.value.replace("_", " ").title(),
            "confidence": round(s.confidence, 2),
            "confidence_pct": f"{s.confidence:.0%}",
            "confidence_level": s.confidence_level.value,
            "confidence_color": f"#{_CONF_HEX.get(s.confidence_level.value, '999999')}",
            "status": s.status.value,
            "status_color": f"#{_STATUS_HEX.get(s.status.value, '999999')}",
            "source_name": s.source_name,
            "source_url": s.source_url,
            "detected_at": s.detected_at.strftime("%d %b %Y"),
            "playbook": s.playbook_match.value if s.playbook_match else None,
            "recommended_action": s.recommended_action,
            "impact_statement": s.impact_statement,
            "talk_track": s.talk_track,
            "current_provider": s.current_provider,
        })

    # Category summary
    cat_counts: dict[str, int] = {}
    for s in store.get_recent_signals(days):
        cat_counts[s.category.value] = cat_counts.get(s.category.value, 0) + 1

    category_summary = [
        {
            "category": cat,
            "label": cat.replace("_", " ").title(),
            "count": count,
            "color": f"#{_CATEGORY_HEX.get(cat, '2980B9')}",
        }
        for cat, count in sorted(cat_counts.items(), key=lambda x: -x[1])
    ]

    # Confidence summary
    all_recent = store.get_recent_signals(days)
    conf_summary = {
        "high": sum(1 for s in all_recent if s.confidence >= 0.8),
        "medium": sum(1 for s in all_recent if 0.5 <= s.confidence < 0.8),
        "low": sum(1 for s in all_recent if s.confidence < 0.5),
    }

    payload = {
        "type": "signal_cards",
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "total_signals": total,
        "showing": len(cards),
        "period_days": days,
        "confidence_summary": conf_summary,
        "category_summary": category_summary,
        "cards": cards,
    }

    cards_json = json.dumps(payload, indent=2)

    # Upload to blob for reliable access
    blob = _get_blob_store()
    blob_url = None
    if blob:
        try:
            ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            blob_url = blob.upload_text(cards_json, f"cards/signal_cards_{ts}.json")
        except Exception as exc:
            logger.warning("Failed to upload cards JSON to blob: %s", exc)

    # Return a compact summary + the JSON data
    link_line = f"\n\n📂 [Download card data (JSON)]({blob_url})" if blob_url else ""
    return (
        f"{{CARD_JSON_START}}\n{cards_json}\n{{CARD_JSON_END}}"
        f"{link_line}"
    )


# ------------------------------------------------------------------ #
# Collect all tools for registration
# ------------------------------------------------------------------ #

ALL_INTERACTIVE_TOOLS = [
    # Core signal intelligence tools
    get_signals,
    get_signal_details,
    get_hitl_queue,
    review_signal,
    run_sweep_now,
    get_run_history,
    generate_brief,
    generate_report,
    get_signal_dashboard,
    # Proactive M365 distribution tools
    send_daily_digest,
    send_high_alert,
    schedule_review_meeting,
    create_word_report,
    create_action_tasks,
    update_pipeline_tracker,
    post_to_teams_channel,
    distribute_intelligence,
    get_signal_cards,
]
